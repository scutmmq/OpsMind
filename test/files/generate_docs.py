"""生成 OpsMind 验收测试用的 DOCX 和 PDF 知识文档。

依赖：python-docx, fpdf2
运行：python generate_docs.py
"""

from pathlib import Path
import json

# ============================================================
# 文档内容定义
# ============================================================

DOCUMENTS = {
    "docx": {
        "邮件系统使用规范": {
            "title": "企业邮件系统使用规范",
            "paragraphs": [
                ("heading1", "企业邮件系统使用规范"),
                ("para", "本文档规定公司邮件系统的使用标准，适用于全体员工。邮箱是企业对内对外沟通的主要工具，"
                         "所有员工必须遵守本规范。"),
                ("heading2", "一、邮箱基本配置"),
                ("para", "公司使用 Microsoft Exchange Online（Office 365）作为邮件平台。每位员工拥有独立的邮箱账号，"
                         "格式为「姓名拼音@company.com」。邮箱容量默认 50GB，单封邮件最大 25MB（含附件）。"),
                ("heading2", "1.1 桌面客户端配置"),
                ("para", "推荐使用 Microsoft Outlook 桌面客户端。首次配置步骤："),
                ("bullet", "打开 Outlook → 文件 → 添加账户"),
                ("bullet", "输入公司邮箱地址（如 zhangsan@company.com）"),
                ("bullet", "选择「Exchange」账户类型（自动发现服务器设置）"),
                ("bullet", "输入域账号密码完成验证"),
                ("bullet", "等待邮箱数据同步完成（首次约需 5-15 分钟）"),
                ("heading2", "1.2 移动端配置"),
                ("para", "iOS：设置 → 邮件 → 账户 → 添加账户 → Exchange → 输入邮箱和密码。"),
                ("para", "Android：下载 Outlook for Android → 添加账户 → 输入公司邮箱 → 按提示完成验证。"),
                ("heading2", "二、邮件使用规范"),
                ("heading3", "2.1 邮件撰写要求"),
                ("bullet", "主题行：简明扼要，格式为「[类别] 主题 - 发件人部门」（如「[审批] Q2 出差申请 - 研发部」）"),
                ("bullet", "收件人：TO=需处理/回复的人，CC=需知悉的人（避免过度抄送）"),
                ("bullet", "正文：使用正式语体，分段清晰。紧急事项可标注【紧急】但不得滥用"),
                ("bullet", "附件：超过 10MB 的附件使用 OneDrive 链接替代，多个附件请打包为 ZIP"),
                ("bullet", "签名：统一使用公司标准签名模板（含姓名、职位、部门、电话、公司 Logo）"),
                ("heading3", "2.2 邮件分类与归档"),
                ("para", "建议使用以下文件夹分类管理邮件："),
                ("bullet", "「01-待处理」：需要回复或处理的邮件"),
                ("bullet", "「02-进行中」：已回复但事项未完结的邮件"),
                ("bullet", "「03-已完结」：事项已完成的归档邮件"),
                ("bullet", "「04-参考」：公告、通知等不需要处理的邮件"),
                ("bullet", "「05-项目/」：按项目名称创建子文件夹归档"),
                ("heading2", "三、安全规范"),
                ("heading3", "3.1 钓鱼邮件识别"),
                ("para", "钓鱼邮件常见特征："),
                ("bullet", "发件人地址伪造（如 admin@company-secure.com 而非 admin@company.com）"),
                ("bullet", "制造紧迫感（「您的账号将在 24 小时内被冻结，请立即验证」）"),
                ("bullet", "索取密码或敏感信息（公司 IT 绝不会通过邮件索取密码）"),
                ("bullet", "附件为可执行文件（.exe、.js、.vbs 等）"),
                ("bullet", "链接指向非公司域名（悬停查看真实 URL）"),
                ("para", "收到可疑邮件请通过 Outlook 的「报告钓鱼」按钮上报，切勿点击链接或下载附件。"),
                ("heading3", "3.2 机密信息发送"),
                ("bullet", "机密级及以上信息不得通过邮件明文发送"),
                ("bullet", "如需发送敏感文件，使用 Office 365 的「加密邮件」功能"),
                ("bullet", "外部收件人（非公司域）发送敏感信息前需获得部门负责人批准"),
                ("heading2", "四、常见问题"),
                ("heading3", "4.1 邮箱无法收发邮件"),
                ("bullet", "检查网络连接是否正常"),
                ("bullet", "检查 Outlook 右下角状态是否为「已连接到 Microsoft Exchange」"),
                ("bullet", "尝试网页版邮箱 https://mail.company.com 确认是否为客户端问题"),
                ("bullet", "网页版正常但客户端异常：控制面板 → 邮件 → 重建配置文件"),
                ("bullet", "均不可用：联系 IT 服务台（分机 8888）"),
                ("heading3", "4.2 邮箱存储空间不足"),
                ("bullet", "清理「已删除邮件」和「垃圾邮件」文件夹"),
                ("bullet", "按附件大小排序，删除含大附件的旧邮件（附件已保存到 OneDrive 的可以删除邮件）"),
                ("bullet", "使用 Outlook 的「邮箱清理」工具自动归档 6 个月前的邮件"),
                ("bullet", "如需扩容至 100GB，提交 IT 设备申请单（需部门负责人审批）"),
                ("heading3", "4.3 自动回复设置"),
                ("para", "休假或出差前请设置自动回复："),
                ("bullet", "Outlook → 文件 → 自动回复 → 设置时间段和回复内容"),
                ("bullet", "回复内容需包含：休假时间、紧急联系人及联系方式"),
                ("bullet", "可选择是否对外部发件人也发送自动回复"),
                ("heading2", "五、违规处罚"),
                ("para", "违反本规范的行为将视情节严重程度给予相应处理："),
                ("bullet", "首次违规：口头提醒并记录"),
                ("bullet", "再次违规：书面警告，抄送部门负责人"),
                ("bullet", "严重违规（如泄露机密信息）：按公司《信息安全管理制度》处理，"
                           "最高可至解除劳动合同"),
                ("para", ""),
                ("para", "文档版本：v2.0 | 最后更新：2026-04-18 | 维护部门：IT 基础设施组"),
            ],
        },
        "信息安全管理制度": {
            "title": "信息安全管理制度（摘要）",
            "paragraphs": [
                ("heading1", "信息安全管理制度（摘要）"),
                ("para", "本制度依据 ISO 27001 信息安全管理体系制定，适用于公司全体员工（含外包和实习人员）。"
                         "全文见公司 Confluence 安全合规空间。"),
                ("heading2", "一、信息资产分类"),
                ("heading3", "1.1 分类标准"),
                ("bullet", "公开：对外发布的信息（官网、公告、产品手册等）。无需特殊保护。"),
                ("bullet", "内部：公司内部流转的信息（项目文档、会议纪要、制度文件等）。泄露可能导致轻微不利影响。"),
                ("bullet", "机密：部门级敏感信息（财务报表、人事档案、客户合同等）。泄露可能导致重大损失。"),
                ("bullet", "绝密：公司级战略信息（商业计划、并购方案、核心算法等）。泄露可能导致生存危机。"),
                ("heading3", "1.2 存储和传输要求"),
                ("para", "不同级别信息的存储和传输要求如下表："),
                ("para", "公开/内部：可使用公司标准设备存储，可通过邮件/IM 内部分享。"),
                ("para", "机密：必须加密存储（BitLocker 全盘加密），仅限「需要知晓」的人员访问。"
                         "传输须通过公司批准的加密通道。"),
                ("para", "绝密：须使用专用加密设备（硬件加密），访问需双人审批+审计日志。"
                         "原则上禁止网络传输。"),
                ("heading2", "二、密码管理"),
                ("bullet", "密码长度不少于 8 位，须含大写字母、小写字母和数字"),
                ("bullet", "禁止使用以下弱密码：123456、password、admin、公司名+年份、工号"),
                ("bullet", "不同系统使用不同密码（禁止一码多用）"),
                ("bullet", "密码 90 天强制更换，不得与之前 5 次密码重复"),
                ("bullet", "推荐使用密码管理器（公司提供 1Password Business 授权）"),
                ("bullet", "发现密码泄露应立即修改并报告信息安全组"),
                ("heading2", "三、设备安全"),
                ("heading3", "3.1 公司配发设备"),
                ("bullet", "启用全盘加密（Windows BitLocker / macOS FileVault）"),
                ("bullet", "安装公司统一的防病毒软件（CrowdStrike Falcon），不得卸载或禁用"),
                ("bullet", "开启屏幕保护程序，设置 5 分钟无操作自动锁定"),
                ("bullet", "系统补丁由 IT 统一推送，员工需在收到通知后 48 小时内完成更新"),
                ("bullet", "离职或换岗时必须交还设备，IT 将执行安全擦除"),
                ("heading3", "3.2 移动设备管理（MDM）"),
                ("bullet", "接入公司邮件/IM 的移动设备须注册 Intune MDM"),
                ("bullet", "设备须设置至少 6 位数字密码或生物识别锁"),
                ("bullet", "丢失或被盗应立即远程擦除公司数据（通过 Intune 自助或联系 IT）"),
                ("heading2", "四、网络安全"),
                ("bullet", "办公网络仅限公司设备接入。个人设备使用 Guest-WiFi（仅外网，无内网权限）"),
                ("bullet", "远程办公必须通过公司 VPN 接入内网（禁止端口映射等绕过方式）"),
                ("bullet", "禁止在公共场所（咖啡厅、机场等）通过不加密 WiFi 处理机密级及以上信息"),
                ("bullet", "禁止私自搭建内网服务（如自建 WiFi 热点、代理服务器等）"),
                ("heading2", "五、数据保护"),
                ("heading3", "5.1 数据备份"),
                ("bullet", "重要工作文件必须同步至 OneDrive for Business（已自动配置）"),
                ("bullet", "代码须提交至公司 GitLab，禁止仅保存在本地"),
                ("bullet", "数据库备份由 IT 统一管理，每日自动备份保留 30 天"),
                ("heading3", "5.2 数据销毁"),
                ("bullet", "纸质机密文件须使用碎纸机销毁（交叉切割型）"),
                ("bullet", "电子存储介质（硬盘、U 盘）报废前须使用 DoD 5220.22-M 标准擦除或物理销毁"),
                ("bullet", "离职员工的邮箱和 OneDrive 数据保留 30 天后自动清除"),
                ("heading2", "六、安全事件响应"),
                ("para", "发现以下情况应立即报告信息安全组（分机 1111，邮箱 security@company.com）："),
                ("bullet", "电脑中病毒/勒索软件"),
                ("bullet", "收到可疑的钓鱼邮件并已点击链接/下载附件"),
                ("bullet", "发现非授权人员在公司区域活动"),
                ("bullet", "设备或存储介质丢失/被盗"),
                ("bullet", "发现系统漏洞或数据泄露"),
                ("para", "安全事件响应流程："),
                ("bullet", "1. 立即断网（拔网线/关 WiFi）以防扩散"),
                ("bullet", "2. 报告信息安全组，保留现场"),
                ("bullet", "3. 信息安全组启动应急响应（2 小时内给出初步评估）"),
                ("bullet", "4. 重大事件 24 小时内上报管理层"),
                ("para", ""),
                ("para", "文档版本：v3.1 | 最后更新：2026-05-20 | 维护部门：信息安全组"),
            ],
        },
    },
}

# ============================================================
# DOCX 生成
# ============================================================

def generate_docx(output_dir: Path):
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for name, doc_def in DOCUMENTS["docx"].items():
        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Microsoft YaHei"
        font.size = Pt(11)

        for ptype, text in doc_def["paragraphs"]:
            if ptype == "heading1":
                h = doc.add_heading(text, level=1)
            elif ptype == "heading2":
                h = doc.add_heading(text, level=2)
            elif ptype == "heading3":
                h = doc.add_heading(text, level=3)
            elif ptype == "bullet":
                doc.add_paragraph(text, style="List Bullet")
            elif ptype == "para":
                if text:
                    doc.add_paragraph(text)

        filepath = output_dir / f"{name}.docx"
        doc.save(str(filepath))
        print(f"  OK {filepath.name}")


# ============================================================
# PDF 生成
# ============================================================

def generate_pdf(output_dir: Path):
    from fpdf import FPDF

    # 注册中文字体 - 使用系统自带字体或回退到内置
    import sys
    font_paths = [
        # Windows
        Path("C:/Windows/Fonts/msyh.ttc"),       # 微软雅黑
        Path("C:/Windows/Fonts/simsun.ttc"),      # 宋体
        Path("C:/Windows/Fonts/simhei.ttf"),      # 黑体
    ]

    cn_font = None
    for fp in font_paths:
        if fp.exists():
            cn_font = str(fp)
            print(f"  -> Using font: {fp.name}")
            break

    if cn_font is None:
        print("  WARNING: No Chinese font found for PDF")
        print("    Copy C:/Windows/Fonts/msyh.ttc to current dir if needed")

    # 合并所有 docx 内容生成两个 PDF
    for name, doc_def in DOCUMENTS["docx"].items():
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)

        if cn_font:
            pdf.add_font("CN", "", cn_font, uni=True)
            pdf.add_font("CN", "B", cn_font, uni=True)

        pdf.add_page()

        for ptype, text in doc_def["paragraphs"]:
            if not text:
                continue

            if cn_font:
                if ptype == "heading1":
                    pdf.set_font("CN", "B", 18)
                    pdf.cell(0, 12, text, new_x="LMARGIN", new_y="NEXT")
                    pdf.ln(6)
                elif ptype == "heading2":
                    pdf.set_font("CN", "B", 14)
                    pdf.cell(0, 10, text, new_x="LMARGIN", new_y="NEXT")
                    pdf.ln(4)
                elif ptype == "heading3":
                    pdf.set_font("CN", "B", 12)
                    pdf.cell(0, 8, text, new_x="LMARGIN", new_y="NEXT")
                    pdf.ln(2)
                elif ptype == "bullet":
                    pdf.set_font("CN", "", 10)
                    pdf.cell(5, 6, "•")
                    pdf.multi_cell(0, 6, text, new_x="LMARGIN", new_y="NEXT")
                else:
                    pdf.set_font("CN", "", 10)
                    pdf.multi_cell(0, 6, text, new_x="LMARGIN", new_y="NEXT")
                    pdf.ln(1)

        filepath = output_dir / f"{name}.pdf"
        pdf.output(str(filepath))
        print(f"  OK {filepath.name}")


# ============================================================
# Main
# ============================================================

def main():
    base = Path(__file__).parent

    docx_dir = base / "docx"
    pdf_dir = base / "pdf"
    docx_dir.mkdir(exist_ok=True)
    pdf_dir.mkdir(exist_ok=True)

    print("Generating DOCX files...")
    generate_docx(docx_dir)

    print()
    print("Generating PDF files...")
    generate_pdf(pdf_dir)

    print()
    print("All test documents generated!")
    print(f"  DOCX: {docx_dir}")
    print(f"  PDF:  {pdf_dir}")


if __name__ == "__main__":
    main()
